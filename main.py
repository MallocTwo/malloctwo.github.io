import os
import pandas as pd
import datetime
from collections import Counter
from flask import Flask, render_template, request, redirect, url_for, send_file, flash
from werkzeug.utils import secure_filename
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import zipfile

app = Flask(__name__)
app.secret_key = 'supersecretkey'  # Needed for flashing messages
app.config['UPLOAD_FOLDER'] = '/Users/ashaw/Desktop/Code/TestWebService/uploads'
app.config['PROCESSED_FOLDER'] = '/Users/ashaw/Desktop/Code/TestWebService/processed'
app.config['ALLOWED_EXTENSIONS'] = {'csv'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in app.config['ALLOWED_EXTENSIONS']

def process_csv(file_path):
    # Today's Date
    today = datetime.date.today()
    formatted_date = today.strftime('%m/%d/%Y')

    # Read CampOps Sheet
    campOp = pd.read_csv(file_path)

    # Remove Total Passengers
    campOp.drop([len(campOp)-1], inplace=True)

    # Remove specific columns
    campOp.drop(['Arrival', 'Food Service', 'Show Time', 'Bake', 'Camp', 'Max', 'Pax', 'Notes', 'Lock'], axis=1, inplace=True)

    # Sort CampOp CSV by time
    timeUnsorted = []
    for findTime in campOp['Tour']:
        convertToString = str(findTime)
        timeInString = convertToString[:5]
        removeSemicolon = timeInString.replace(":", "")
        timeUnsorted.append(int(removeSemicolon))
    campOp['Time'] = timeUnsorted
    campOp.sort_values(by='Time', inplace=True)
    campOp.drop('Time', axis=1, inplace=True)
    campOp.reset_index(drop=True, inplace=True)

    # Create DataFrame with Ship, Time, and Tour
    shipTimeTour = pd.DataFrame(columns=['Ship', 'Time', 'Tour'])
    shipCounter = Counter(campOp['Ship'])
    j = 0
    for shipName, shipCount in shipCounter.items():
        for i in range(0, len(campOp['Ship'])):
            if str(shipName) == str(campOp['Ship'][i]):
                getTourName = str(campOp['Tour'][i])
                fixedTourName = ' - '.join(part.strip() for part in getTourName.split('-'))
                getTour = fixedTourName[11:]
                getTime = fixedTourName[:8]
                shipTimeTour.loc[j, 'Ship'] = str(shipName)
                shipTimeTour.loc[j, 'Time'] = str(getTime)
                shipTimeTour.loc[j, 'Tour'] = str(getTour)
                j += 1

    # Create Each Ship's Settlement Documents
    new_ship_time_tour = Counter(shipTimeTour['Ship'])
    index = [0]
    processed_files = []
    for x, y in new_ship_time_tour.items():
        index.append(int(y))
        doc = Document()
        table = doc.add_table(rows=3+y, cols=6)
        table.style = 'Table Grid'
        for row in range(2):
            a = table.cell(row, 0)
            b = table.cell(row, 5)
            merged_cell = a.merge(b)
        table.cell(0, 0).text = f"Date: {formatted_date}"
        table.cell(1, 0).text = f"Ship: {x}"
        table.cell(2, 0).text = "Tour"
        table.cell(2, 1).text = "Time"
        table.cell(2, 2).text = "Adult"
        table.cell(2, 3).text = "Child"
        table.cell(2, 4).text = "Comp"
        table.cell(2, 5).text = "Other"
        start = 0
        end = 0
        if len(index) == 2:
            start = index[0]
            end = index[1]
        if len(index) > 2:
            for i in range(len(index)-1):
                start += index[i]
            for i in range(len(index)):
                end += index[i]
        count = 0
        for tour in range(start, end):
            table.cell(3 + count, 0).text = shipTimeTour['Tour'][tour]
            count += 1
        count = 0
        for tour in range(start, end):
            table.cell(3 + count, 1).text = shipTimeTour['Time'][tour]
            count += 1
        processed_filename = os.path.join(app.config['PROCESSED_FOLDER'], f'{x}.docx')
        doc.save(processed_filename)
        processed_files.append(processed_filename)

        #Create Transport Company Settlement
        tourCo = pd.DataFrame(new_ship_time_tour.items(), columns=['Ship', 'Number'])
        transportCompany = pd.DataFrame(columns=['Transport', 'Ship', 'Time', 'Tour'])
        tourCo.drop([len(tourCo)-1], inplace= True)
        for findTourCo in tourCo['Ship']:
            count = 0
            for trans in campOp['Ship']:
                if findTourCo == trans and 'ATA' != campOp['Transport'][count]:
                    getTourName = str(campOp['Tour'][count])
                    fixedTourName = ' - '.join(part.strip() for part in getTourName.split('-'))
                    getTour = fixedTourName[11:]
                    getTime = fixedTourName[:8]
                    transportCompany.loc[count, 'Transport'] = campOp['Transport'][count]
                    transportCompany.loc[count, 'Ship'] = trans
                    transportCompany.loc[count, 'Time'] = str(getTime)
                    transportCompany.loc[count, 'Tour'] = str(getTour)
                count += 1
        transportCompany.reset_index(drop=True, inplace=True)
        transportCounter = Counter(transportCompany['Transport'])
        act = transportCompany.iloc[:transportCounter['ACT']]
        act.reset_index(drop=True, inplace=True)
        hap = transportCompany.iloc[transportCounter['ACT']:]
        hap.reset_index(drop=True, inplace=True)

        #ACT Settlement Sheet
        actCounter = Counter(act['Ship'])
        act_index = [0]
        actdoc = Document()
        title = actdoc.add_heading("ACT", level=1)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for x, y in actCounter.items():
            act_index.append(int(y))
            table = actdoc.add_table(rows=3+y, cols=6)
            table.style = 'Table Grid'
            for row in range(2):
                a = table.cell(row, 0)
                b = table.cell(row, 5)
                merged_cell = a.merge(b)
            table.cell(0, 0).text = f"Date: {formatted_date}"
            table.cell(1,0).text = f"Ship: {x}"
            table.cell(2, 0).text = "Tour"
            table.cell(2, 1).text = "Time"
            table.cell(2, 2).text = "Adult"
            table.cell(2, 3).text = "Child"
            table.cell(2, 4).text = "Comp"
            table.cell(2, 5).text = "Other"
            start = 0
            end = 0
            if len(act_index) == 2:
                start = act_index[0]
                end = act_index[1]
            if len(act_index) > 2:
                for i in range(len(act_index)-1):
                    start += act_index[i]
                for i in range(len(act_index)):
                    end += act_index[i]
            count = 0
            for tour in range(start, end):
                table.cell(3 + count, 0).text = act['Tour'][tour]
                count += 1
            count = 0
            for tour in range(start, end):
                table.cell(3 + count, 1).text = act['Time'][tour]
                count += 1
            actdoc.add_paragraph()
        if len(actCounter) != 0:
            processed_filename = os.path.join(app.config['PROCESSED_FOLDER'], 'ACT.docx')
            actdoc.save(processed_filename)
            processed_files.append(processed_filename)

        #HAP Settlement Sheet
        hapCounter = Counter(hap['Ship'])
        hap_index = [0]
        hapdoc = Document()
        title = hapdoc.add_heading("HAP", level=1)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for x, y in hapCounter.items():
            hap_index.append(int(y))
            table = hapdoc.add_table(rows=3+y, cols=6)
            table.style = 'Table Grid'
            for row in range(2):
                a = table.cell(row, 0)
                b = table.cell(row, 5)
                merged_cell = a.merge(b)
            table.cell(0, 0).text = f"Date: {formatted_date}"
            table.cell(1,0).text = f"Ship: {x}"
            table.cell(2, 0).text = "Tour"
            table.cell(2, 1).text = "Time"
            table.cell(2, 2).text = "Adult"
            table.cell(2, 3).text = "Child"
            table.cell(2, 4).text = "Comp"
            table.cell(2, 5).text = "Other"
            start = 0
            end = 0
            if len(hap_index) == 2:
                start = hap_index[0]
                end = hap_index[1]
            if len(hap_index) > 2:
                for i in range(len(hap_index)-1):
                    start += hap_index[i]
                for i in range(len(hap_index)):
                    end += hap_index[i]
            count = 0
            for tour in range(start, end):
                table.cell(3 + count, 0).text = hap['Tour'][tour]
                count += 1
            count = 0
            for tour in range(start, end):
                table.cell(3 + count, 1).text = hap['Time'][tour]
                count += 1
            hapdoc.add_paragraph()
        if len(hapCounter) != 0:
            processed_filename = os.path.join(app.config['PROCESSED_FOLDER'], f'HAP.docx')
            hapdoc.save(processed_filename)
            processed_files.append(processed_filename)

    # Zip all the processed files
    zip_filename = os.path.join(app.config['PROCESSED_FOLDER'], 'Settlement.zip')
    with zipfile.ZipFile(zip_filename, 'w') as zipf:
        for file in processed_files:
            zipf.write(file, os.path.basename(file))

    return zip_filename

@app.route('/')
def home():
    return render_template('test.html')

@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        flash('No file part')
        return redirect(request.url)
    file = request.files['file']
    if file.filename == '':
        flash('No selected file')
        return redirect(request.url)
    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(file_path)
        processed_file_path = process_csv(file_path)
        return redirect(url_for('download_file', filename=os.path.basename(processed_file_path)))
    else:
        flash('File extension not allowed. Please upload a CSV file.')
        return redirect(request.url)

@app.route('/download/<filename>')
def download_file(filename):
    return send_file(os.path.join(app.config['PROCESSED_FOLDER'], filename), as_attachment=True)

if __name__ == '__main__':
    os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
    os.makedirs(app.config['PROCESSED_FOLDER'], exist_ok=True)
    app.run(host='127.0.0.1', port=9000)
